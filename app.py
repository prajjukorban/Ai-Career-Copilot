import streamlit as st
from pypdf import PdfReader
import google.generativeai as genai
import time
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

genai.configure(api_key=env.API_KEY)
model = genai.GenerativeModel("gemini-2.5-flash")

st.set_page_config(page_title="AI Career Copilot", layout="centered")

st.markdown("""
    <style>
        .main { background-color: #0e1117; }
        .title { text-align: center; font-size: 40px; font-weight: bold; color: #4CAF50; }
        .subtitle { text-align: center; color: #bbbbbb; margin-bottom: 25px; }
        .output-box {
            padding: 20px; border-radius: 12px;
            background-color: #1f2937; color: #e5e7eb;
            font-size: 16px; white-space: pre-wrap;
        }
        .section-header { color: #4CAF50; font-size: 18px; font-weight: 600; margin-top: 10px; }
    </style>
""", unsafe_allow_html=True)

st.markdown('<div class="title">🚀 AI Career Copilot</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">Upload your resume OR fill in your profile → Role Prediction, ATS Score & Skill Gaps</div>', unsafe_allow_html=True)

mode = st.radio(
    "How would you like to proceed?",
    options=["📄 Upload My Resume", "✍️ Don't Have a Resume (Fill Manually)"],
    horizontal=True
)
st.divider()

resume_text = ""
manual_profile = {}

# ── DOCX helpers ──────────────────────────────────────────────────────────────
def add_hr(doc, color="2E8B57"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def section_heading(doc, text):
    h = doc.add_paragraph(text)
    h.runs[0].bold = True
    h.runs[0].font.size = Pt(11)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x73, 0x48)
    h.paragraph_format.space_before = Pt(10)
    h.paragraph_format.space_after = Pt(3)

def build_resume_docx(profile: dict) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(profile.get("name") or "Your Name")
    r.bold = True; r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x1A, 0x73, 0x48)

    # Contact line
    parts = []
    if profile.get("location"):   parts.append(f"📍 {profile['location']}")
    if profile.get("field"):      parts.append(f"🎯 {profile['field']}")
    if profile.get("experience"): parts.append(f"⏳ {profile['experience']}")
    if parts:
        cp = doc.add_paragraph(" | ".join(parts))
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cp.runs:
            r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x55,0x55,0x55)

    add_hr(doc)

    # Summary
    if profile.get("about"):
        section_heading(doc, "PROFESSIONAL SUMMARY")
        p = doc.add_paragraph(profile["about"])
        for r in p.runs: r.font.size = Pt(10.5)
        add_hr(doc)

    # Skills
    if profile.get("skills"):
        section_heading(doc, "SKILLS")
        p = doc.add_paragraph(", ".join(profile["skills"]))
        for r in p.runs: r.font.size = Pt(10.5)
        add_hr(doc)

    # Education
    if profile.get("education"):
        section_heading(doc, "EDUCATION")
        p = doc.add_paragraph()
        r = p.add_run(profile["education"]); r.bold = True; r.font.size = Pt(10.5)
        if profile.get("field"):
            p.add_run(f"  |  {profile['field']}").font.size = Pt(10.5)
        add_hr(doc)

    # Projects
    if profile.get("projects", "").strip():
        section_heading(doc, "PROJECTS")
        for line in profile["projects"].strip().splitlines():
            line = line.strip().lstrip("-• ").strip()
            if line:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.add_run(line).font.size = Pt(10.5)
        add_hr(doc)

    # Certifications
    if profile.get("certifications", "").strip():
        section_heading(doc, "CERTIFICATIONS")
        for line in profile["certifications"].strip().splitlines():
            line = line.strip().lstrip("-• ").strip()
            if line:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.add_run(line).font.size = Pt(10.5)

    # Desired role footer
    if profile.get("desired_role"):
        doc.add_paragraph()
        p = doc.add_paragraph(f"🎯 Actively seeking: {profile['desired_role']}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(10); r.font.italic = True
            r.font.color.rgb = RGBColor(0x44,0x44,0x44)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

# ── MODE 1: Upload Resume ──────────────────────────────────────────────────────
if mode == "📄 Upload My Resume":
    uploaded_file = st.file_uploader("📂 Upload Resume (PDF only)", type=["pdf"])
    if uploaded_file:
        reader = PdfReader(uploaded_file)
        text = "".join(p.extract_text() or "" for p in reader.pages)
        resume_text = text
        with st.expander("📄 View Extracted Resume Text"):
            st.write(text or "No text could be extracted.")

# ── MODE 2: Manual Entry ───────────────────────────────────────────────────────
else:
    st.markdown('<div class="section-header">👤 Tell us about yourself</div>', unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        name = st.text_input("Full Name", placeholder="e.g. Ravi Kumar")
        experience_years = st.selectbox("Years of Experience",
            ["Fresher (0 years)","1–2 years","3–5 years","6–10 years","10+ years"])
        education = st.selectbox("Highest Education",
            ["High School","Diploma","Bachelor's","Master's","PhD","Other"])
    with col2:
        field = st.selectbox("Field / Domain",
            ["Software Engineering","Data Science / ML / AI","Web Development",
             "Mobile Development","DevOps / Cloud","Cybersecurity",
             "Product Management","UI/UX Design","Business Analyst",
             "Finance / Accounting","Marketing / Growth","HR / Recruitment",
             "Sales","Healthcare","Education","Other"])
        desired_role = st.text_input("Desired Job Role (optional)", placeholder="e.g. Data Analyst")
        location = st.text_input("Preferred Work Location (optional)", placeholder="e.g. Bangalore, Remote")

    st.markdown('<div class="section-header">🛠️ Skills</div>', unsafe_allow_html=True)
    SKILL_OPTIONS = {
        "Programming Languages": ["Python","JavaScript","Java","C++","C#","TypeScript","Go","Rust","PHP","Swift","Kotlin","Ruby","R","Scala","MATLAB"],
        "Web / Frontend": ["HTML","CSS","React","Vue.js","Angular","Next.js","Tailwind CSS","Bootstrap","Redux","GraphQL"],
        "Backend / Frameworks": ["Node.js","Django","Flask","FastAPI","Spring Boot","Express.js","Laravel","Ruby on Rails",".NET"],
        "Data & ML": ["Pandas","NumPy","Scikit-learn","TensorFlow","PyTorch","Keras","SQL","Excel","Power BI","Tableau","Spark"],
        "Cloud & DevOps": ["AWS","Azure","GCP","Docker","Kubernetes","Terraform","CI/CD","Jenkins","Linux","Git / GitHub"],
        "Design & Other": ["Figma","Adobe XD","Photoshop","Illustrator","Canva","SEO","Google Analytics","Agile / Scrum","Jira","Communication"],
    }
    selected_skills = []
    for category, skills in SKILL_OPTIONS.items():
        with st.expander(f"🔹 {category}"):
            chosen = st.multiselect(f"Select from {category}", skills, key=category)
            selected_skills.extend(chosen)

    custom_skills = st.text_input("➕ Add custom skills (comma-separated)", placeholder="e.g. Prompt Engineering, Blender")
    if custom_skills:
        selected_skills.extend([s.strip() for s in custom_skills.split(",") if s.strip()])

    st.markdown('<div class="section-header">📝 Brief About Yourself (optional)</div>', unsafe_allow_html=True)
    about = st.text_area("Short bio / career summary",
        placeholder="e.g. I'm a fresher passionate about backend development...", height=100)

    st.markdown('<div class="section-header">🏆 Projects & Certifications (optional)</div>', unsafe_allow_html=True)
    col3, col4 = st.columns(2)
    with col3:
        projects = st.text_area("Key Projects",
            placeholder="- E-commerce site with React + Node\n- ML model for house price prediction", height=100)
    with col4:
        certifications = st.text_area("Certifications / Courses",
            placeholder="- AWS Certified Cloud Practitioner\n- Google Data Analytics Certificate", height=100)

    manual_profile = dict(name=name, experience=experience_years, education=education,
                          field=field, desired_role=desired_role, location=location,
                          skills=selected_skills, about=about, projects=projects,
                          certifications=certifications)

    resume_text = f"""Name: {name or 'Not provided'}
Field: {field}
Experience: {experience_years}
Education: {education}
Desired Role: {desired_role or 'Not specified'}
Location: {location or 'Not specified'}
Skills: {', '.join(selected_skills) if selected_skills else 'None selected'}
About: {about or 'Not provided'}
Projects:
{projects or 'Not provided'}
Certifications:
{certifications or 'Not provided'}""".strip()

    # ── Generate Resume Section ───────────────────────────────────────────────
    st.divider()
    st.markdown('<div class="section-header">📄 Generate & Download Your Resume</div>', unsafe_allow_html=True)
    st.caption("Builds a clean, professional Word (.docx) resume from the info you filled above.")

    has_basic_info = bool(name or selected_skills or about or projects)

    if not has_basic_info:
        st.info("💡 Fill in at least your name, skills, or a short bio above to generate your resume.")
    else:
        if st.button("📝 Build My Resume (.docx)", use_container_width=True):
            with st.spinner("Generating your resume..."):
                docx_bytes = build_resume_docx(manual_profile)
            file_name = f"{(name or 'resume').replace(' ', '_')}_resume.docx"
            st.success("✅ Resume ready!")
            st.download_button(
                label="⬇️ Download Resume (.docx)",
                data=docx_bytes,
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

# ── Analyze Button ─────────────────────────────────────────────────────────────
st.divider()

if not resume_text.strip():
    st.info("📎 Please upload a resume or fill in your profile above." if mode == "📄 Upload My Resume"
            else "📝 Fill in at least a few fields above to get your analysis.")
else:
    if st.button("✨ Analyze My Profile", use_container_width=True):
        progress = st.progress(0)
        for i in range(100):
            time.sleep(0.01)
            progress.progress(i + 1)

        source_label = "resume" if mode == "📄 Upload My Resume" else "manually entered profile"
        query = f"""
This is a candidate's {source_label}:

{resume_text}

Tasks:
1. Predict the most suitable job role
2. Give an ATS / Profile score out of 100 (be strict and honest)
3. List the top missing skills for that role
4. Give 3–4 actionable improvement tips

Keep each section concise and direct. Be honest like a real career mentor.

Format exactly like this:
Role:
Score:
Missing Skills:
Advice:
"""
        with st.spinner("🤖 AI is analyzing your profile..."):
            result = model.generate_content(query)
            response = result.text.strip()

        st.markdown("### 🔮 Career Insights")
        st.markdown(f'<div class="output-box">{response}</div>', unsafe_allow_html=True)

        st.download_button(
            label="📥 Download Career Report (.txt)",
            data=response,
            file_name="career_report.txt",
            mime="text/plain",
            use_container_width=True
        )
